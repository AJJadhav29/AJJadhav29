# import streamlit as st
# import google.generativeai as genai
# import fitz  # PyMuPDF
# import docx
# import pandas as pd

# # Load API key securely from Streamlit secrets
# api_key = st.secrets["GEMINI_API_KEY"]

# # Configure the Gemini model
# genai.configure(api_key=api_key)
# model = genai.GenerativeModel("gemini-1.5-flash")

# # Initialize chat history
# if "messages" not in st.session_state:
#     st.session_state["messages"] = []

# st.set_page_config(page_title="Tech Career Coach", layout="centered")
# st.title("👩‍💻Tech Career Coach")
# st.markdown("Empowering growth in tech careers through personalized guidance.")

# # One-shot Gemini call for fresh prompts (without chat history)
# def ask_single_prompt(prompt_text):
#     try:
#         response = model.generate_content(prompt_text)
#         return response.text
#     except Exception as e:
#         st.error(f"❌ Error: {str(e)}")
#         return None


# # Input from user
# user_input = st.text_input("💬 Ask me anything about tech careers, mentorship, or STEM growth:")
# st.caption("Examples: 'How do I transition into tech after a career gap?'")
# if user_input:
#     st.session_state["last_input"] = user_input
#     st.session_state["messages"].append({"role": "user", "parts": user_input})
#     try:
#         response = model.generate_content(st.session_state["messages"])
#         st.session_state["messages"].append({"role": "model", "parts": response.text})
#     except Exception as e:
#         st.error(f"❌ Error: {str(e)}")

# # Display chat history
# for msg in st.session_state["messages"]:
#     who = "🧑" if msg["role"] == "user" else "🤖"
#     st.markdown(f"{who}: **{msg['parts']}**")

# # Show latest user input again after actions
# if "last_input" in st.session_state and not user_input:
#     st.markdown(f"🧑: **{st.session_state['last_input']}**")
    
# # Resume File Upload
# def extract_text_from_file(uploaded_file):
#     file_type = uploaded_file.name.split('.')[-1].lower()

#     if file_type == "txt":
#         return uploaded_file.read().decode("utf-8")
#     elif file_type == "pdf":
#         pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
#         text = ""
#         for page in pdf:
#             text += page.get_text()
#         return text
#     elif file_type == "docx":
#         doc = docx.Document(uploaded_file)
#         text = "\n".join([para.text for para in doc.paragraphs])
#         return text
#     else:
#         return "Unsupported file type."

# uploaded_file = st.file_uploader("📎 Upload your resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

# if uploaded_file:
#     extracted_text = extract_text_from_file(uploaded_file)
#     prompt = f"Please review my resume and provide feedback:\n{extracted_text}"
#     st.markdown(f"🧑: Uploaded resume for review.")
#     reply = ask_single_prompt(prompt)
#     if reply:
#         st.markdown(f"🤖: **{reply}**")

# # ----- 🔹 Mentor Match Suggestions (Static) -----
# st.markdown("### 🧑‍🏫 Find a Mentor")

# mentors_df = pd.read_csv("mentors.csv")

# field_choice = st.selectbox("Choose your tech field for mentorship:", mentors_df["field"].unique())

# filtered = mentors_df[mentors_df["field"] == field_choice]

# if not filtered.empty:
#     mentor = filtered.sample(1).iloc[0]  # Pick one at random, or use .iloc[0] for first match
#     st.markdown(f"""
#     👩 Mentor Match: **{mentor['name']}**  
#     🏷️ Field: {mentor['field']}  
#     🌍 Location: {mentor['location']}  
#     📧 Email: `{mentor['email']}`  
#     💼 Experience: {mentor['experience']} years
#     """)
# else:
#     st.warning("No mentors found for this field.")

# # ----- 🔹 Learning Path Suggestion -----
# st.markdown("### 📚 Personalized Learning Path")

# domain = st.selectbox("Choose a career track:", [
#     "Frontend Development", 
#     "Data Science", 
#     "Cloud Engineering", 
#     "Cybersecurity", 
#     "Product Management", 
#     "AI/ML", 
#     "UX Design", 
#     "Backend Development", 
#     "Full Stack Development", 
#     "Data Engineering"
# ])

# if st.button("Suggest Learning Path"):
#     prompt = f"Suggest a beginner-to-intermediate learning path using freeCodeCamp or Coursera for interested in becoming a {domain}. Include key skills, certifications, and estimated timeline."
#     st.markdown(f"🧑: **{prompt}**")
#     reply = ask_single_prompt(prompt)
#     if reply:
#         st.markdown(f"🤖: **{reply}**")


# # ----- Multilingual Translation -----
# language = st.selectbox("🌐 Respond in:", ["English", "Spanish", "Hindi", "French"])

# if language != "English" and st.session_state["messages"]:
#     last_response = st.session_state["messages"][-1]["parts"]
#     translation_prompt = f"Translate the following text into {language}:\n\n{last_response}"
#     translated = ask_single_prompt(translation_prompt)
#     if translated:
#         st.markdown(f"🌐 **Translated ({language}):** {translated}")



#--------------- NEW Design CODE-------------------------#
import streamlit as st
import google.generativeai as genai
import fitz
import docx
import pandas as pd

# --- PAGE CONFIG ---
st.set_page_config(page_title="Tech Career Coach", layout="wide")

# --- SIDEBAR NAVIGATION ---
st.sidebar.title("🔍 Navigation")
menu = st.sidebar.radio("Go to", ["Chatbot", "Resume Review", "Mentor Match", "Learning Path"])

# --- GEMINI CONFIG ---
api_key = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=api_key)
model = genai.GenerativeModel("gemini-1.5-flash")

# --- SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state["messages"] = []

# --- PROMPT FUNCTION ---
def ask_single_prompt(prompt_text):
    try:
        response = model.generate_content(prompt_text)
        return response.text
    except Exception as e:
        st.error(f"❌ Error: {str(e)}")
        return None

# --- CHATBOT PAGE ---
if menu == "Chatbot":
    st.title("💬 Tech Career Coach Chatbot")
    st.caption("Ask about careers, mentorship, or learning paths.")

    for msg in st.session_state["messages"]:
        who = "🧑" if msg["role"] == "user" else "🤖"
        st.markdown(f"{who}: **{msg['parts']}**")

    user_input = st.chat_input("Type your question...")
    if user_input:
        st.session_state["messages"].append({"role": "user", "parts": user_input})
        try:
            response = model.generate_content(st.session_state["messages"])
            st.session_state["messages"].append({"role": "model", "parts": response.text})
            st.rerun()
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

# --- RESUME REVIEW PAGE ---
elif menu == "Resume Review":
    st.title("📎 Resume Reviewer")
    uploaded_file = st.file_uploader("Upload your resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

    def extract_text_from_file(uploaded_file):
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == "txt":
            return uploaded_file.read().decode("utf-8")
        elif file_type == "pdf":
            pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "".join([page.get_text() for page in pdf])
        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        return "Unsupported file type."

    if uploaded_file:
        extracted_text = extract_text_from_file(uploaded_file)
        st.markdown("**🧑: Uploaded resume for review.**")
        reply = ask_single_prompt(f"Please review this resume:{extracted_text}")
        if reply:
            st.markdown(f"🤖: {reply}")

# --- MENTOR MATCH PAGE ---
elif menu == "Mentor Match":
    st.title("🧑‍🏫 Mentor Match")
    df = pd.read_csv("mentors.csv")

    field = st.selectbox("Field of interest:", sorted(df["field"].unique()))
    gender = st.selectbox("Preferred gender:", ["Any", "Female", "Male"])
    location = st.selectbox("Preferred location:", ["Any"] + sorted(df["location"].unique()))

    filtered = df[df["field"] == field]
    if gender != "Any":
        filtered = filtered[filtered["gender"] == gender]
    if location != "Any":
        filtered = filtered[filtered["location"] == location]

    st.subheader("🎯 Your Mentor Match")
    if not filtered.empty:
        for _, mentor in filtered.head(3).iterrows():
            st.markdown(f"""
            **👤 {mentor['name']}**  
            📍 {mentor['location']} | 💼 {mentor['field']} | 🚻 {mentor['gender']}  
            📧 `{mentor['email']}` | 🕓 {mentor['experience']} yrs experience
            ---
            """)
    else:
        st.warning("No mentors found. Try changing your filters.")

# --- LEARNING PATH PAGE ---
elif menu == "Learning Path":
    st.title("📚 Personalized Learning Path")
    domain = st.selectbox("Choose a career track:", [
        "Frontend Development", "Data Science", "Cloud Engineering",
        "Cybersecurity", "Product Management", "AI/ML",
        "UX Design", "Backend Development", "Full Stack Development", "Data Engineering"
    ])
    if st.button("Suggest Learning Path"):
        prompt = f"Suggest a beginner-to-intermediate learning path using freeCodeCamp or Coursera for someone interested in becoming a {domain}. Include key skills and timeline."
        reply = ask_single_prompt(prompt)
        if reply:
            st.markdown(f"🤖: {reply}")










# # 📚 Recommend Courses Button
# if st.button("📚 Recommend Courses for Data Science"):
#     course_prompt = "Recommend beginner-friendly Coursera or edX courses for women who want to start a career in Data Science."
#     try:
#         response = model.generate_content(course_prompt)
#         st.markdown(f"🤖: **{response.text}**")
#     except Exception as e:
#         st.error(f"❌ Error: {str(e)}")

# # 📄 Resume Review Section
# st.markdown("### 📄 Resume / Cover Letter Review")
# resume_text = st.text_area("Paste your resume or cover letter for feedback:")

# if st.button("🔍 Review My Resume"):
#     resume_prompt = f"Review this resume or cover letter and suggest improvements:\n\n{resume_text}"
#     try:
#         response = model.generate_content(resume_prompt)
#         st.markdown(f"🤖: **{response.text}**")
#     except Exception as e:
#         st.error(f"❌ Error: {str(e)}")



